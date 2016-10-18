# -*- coding:utf-8 -*-

import os, io, glob, re
import time
import json
from django.http import JsonResponse, HttpResponse, Http404
from django.db.models.deletion import ProtectedError
from django.db.models import Q
from django.conf import settings
from django.utils import timezone
from django.core.paginator import Paginator
from django.utils.translation import ugettext as _
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from index.models import CartridgeItemName, CartridgeType, CartridgeItem, City
from index.helpers import check_ajax_auth
from docs.models import RefillingCart, SCDoc
from docs.helpers import group_names, localize_date
from service.helpers import SevercartConfigs
from events.helpers import do_timezone


import logging
logger = logging.getLogger(__name__)

@check_ajax_auth
def del_cart_name(request):
    """Удаляем имя расходного материала.
    """
    if request.method != 'POST':
        return HttpResponse('<h1>' + _('Only use POST requests!') + '</h1>')
    
    resp_dict = dict()
    cart_name_id = request.POST.get('cart_name_id', '')
    atype = request.POST.get('atype', '')

    try:
        cart_name_id = int(cart_name_id)
    except ValueError:
        cart_name_id = 0

    if atype == 'cart_name':
        try:
            m1 = CartridgeItemName.objects.get(pk=cart_name_id)
        except CartridgeItemName.DoesNotExist:
            resp_dict['error'] = '1'
            resp_dict['text']  = _('The object with the ID is not found.')
        try:
            m1.delete()
        except ProtectedError:
            resp_dict['error'] = '1'
            resp_dict['text']  = _('Name can not be removed, ie. other objects reference it.')
        else:
            resp_dict['error'] = '0'
            resp_dict['text']  = _('Name deleted successfully.')
    elif atype == 'cart_type':
        try:
            m1 = CartridgeType.objects.get(pk=cart_name_id)
        except CartridgeType.DoesNotExist:
            resp_dict['error'] = '1'
            resp_dict['text']  = _('The object with the ID is not found.')
        try:
            m1.delete()
        except ProtectedError:
            resp_dict['error'] = '1'
            resp_dict['text']  = _('Type can not be deleted, ie the other sites link to it.')
        else:
            resp_dict['error'] = '0'
            resp_dict['text']  = _('Name deleted successfully.')
    return JsonResponse(resp_dict, safe=False)


@check_ajax_auth
def del_city(request):
    """Удаление записи о городе.
    """
    ansver = dict()
    if request.method != 'POST':
        return HttpResponse('<h1>' + _('Only use POST requests!') + '</h1>')
    
    citi_id = request.POST.get('select', 0)

    try:
        citi_id = int(citi_id)
    except ValueError:
        citi_id = 0
    try:
        m1 = City.objects.get(pk=citi_id)
    except City.DoesNotExist:
        ansver['error'] = '1'
        ansver['text']  = _('City object not found.')
        return JsonResponse(ansver)

    try:
        m1.delete()
    except ProtectedError:
        ansver['error'] = '1'
        ansver['text']  = _('City can not be deleted, ie the other firms link to it.')
    else:
        ansver['error'] = '0'
        ansver['text']  = _('City deleted successfully.')
    return JsonResponse(ansver)


@check_ajax_auth
def generate_act(request):
    """Генерация нового docx документа. Вью возвращает Url с свежезгенерированным файлом. 
    """
    resp_dict = dict()
    if request.method != 'POST':
        return HttpResponse('<h1>' + _('Only use POST requests!') + '</h1>')
    # использование глобальных переменных не очень хороший приём
    # но он позволяет упростить программный код
    total_pages_count = 0

    def add_footer(page_num=1):
        # форматирование нижнего колонтитула копирайтом и количеством страниц
        for i in range(4):
            document.add_paragraph("")

        page_number_string = _('Page %(num)s from %(total)s') % {'num': page_num, 'total': total_pages_count}
        p = document.add_paragraph('severcart.org' + ' '*120 + page_number_string)
        
    doc_id = request.POST.get('doc_id', '')
    doc_action = request.POST.get('doc_action', '')
    pages_count = 0
    try:
        doc_id = int(doc_id)
    except ValueError:
        doc_id = 0

    try:
        m1 = RefillingCart.objects.get(pk=doc_id)
    except RefillingCart.DoesNotExist:
        resp_dict['error'] = '1'
        resp_dict['text']  = _('The object with the ID is not found.')
        return JsonResponse(resp_dict, safe=False)

    jsontext = m1.json_content
    jsontext = json.loads(jsontext)

    co = len(jsontext) # количество передаваемых картриджей на заправку

    if not os.path.exists(settings.STATIC_ROOT_DOCX):
        os.makedirs(settings.STATIC_ROOT_DOCX)

    # ротация файлов
    files = filter(os.path.isfile, glob.glob(settings.STATIC_ROOT_DOCX + '\*.docx'))
    files = list(files)
    files.sort(key=lambda x: os.path.getmtime(x))
    try:
        if len(files) > settings.MAX_COUNT_DOCX_FILES:
            os.remove(files[0])
    except:
        pass
    # производим инициализацию некоторых переменных начальными значениями
    sender_full_name    = request.user.fio
    recipient_full_name = ' '*50

    if not((doc_action == "docx_with_group") or (doc_action == "docx_without_group")):
        # если действие указано не верное, сообщаем об это и прекращаем работу 
        # скрипта
        resp_dict['error'] = '1'
        resp_dict['text']  = _('This action is not implemented.')
        return JsonResponse(resp_dict)
    

    if doc_action == 'docx_with_group':
        docx_file_name = m1.number + '_' + str(request.user.pk) +'_1.docx'
        file_full_name = os.path.join(settings.STATIC_ROOT_DOCX, docx_file_name)
        jsontext = group_names(jsontext)
        header_cell_one = _('The name of the cartridge')
        header_cell_two = _('Amount cartridges')
        
    if doc_action == 'docx_without_group':
        docx_file_name = m1.number + '_' + str(request.user.pk) +'_0.docx'
        file_full_name = os.path.join(settings.STATIC_ROOT_DOCX, docx_file_name)
        header_cell_one = _('Cartridge number')
        header_cell_two = _('The name of the cartridge')

    # генерация печатной версии документа без группировки наименований
    document = Document()
    
    # добавляем шапку для документа
    doc_number = str(m1.number) + '/' + str(request.user.pk)
    act_number_string = _('The act of transferring cartridges # %(doc_number)s from  %(date_created)s') % {'doc_number': doc_number, 'date_created': localize_date(m1.date_created)}
    hh1 = document.add_heading(act_number_string, level=2)
    
    hh2 = document.add_heading(str(request.user.departament), level=2)
    
    document.add_paragraph("") # добавляем оступ сверху
    document.add_paragraph("")

    # рисуем таблицу на первой странице
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = header_cell_one
    hdr_cells[1].text = header_cell_two
    if (len(jsontext) <= settings.MAX_TABLE_ROWS_FIRST_PAGE):
        for item in jsontext:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item[0])
            row_cells[1].text = str(item[1])
        total_pages_count = 1

    if (len(jsontext) > settings.MAX_TABLE_ROWS_FIRST_PAGE):
        first_part  = jsontext[0:settings.MAX_TABLE_ROWS_FIRST_PAGE-1]
        second_part = jsontext[settings.MAX_TABLE_ROWS_FIRST_PAGE:]
        
        p = Paginator(second_part, settings.MAX_TABLE_ROWS)
        total_pages_count = 1 + p.num_pages

        for item in first_part:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item[0])
            row_cells[1].text = str(item[1])

        pages_count += 1    
        add_footer(pages_count)
        document.add_page_break()
        # далее с каждой новой страницы рисуем новую таблицу с 
        # продолжением печати данных
        for pg in range(p.num_pages):
            stranica = p.page(pg + 1)
            # на каждой новой странице печатаем заново новый заголовок
            table = document.add_table(rows=1, cols=2, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = header_cell_one
            hdr_cells[1].text = header_cell_two
            for item in stranica:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item[0])
                row_cells[1].text = str(item[1])

            # если страница последняя то разрыв страницы не добавляем
            if  pg != (p.num_pages - 1):
                pages_count += 1
                document.add_page_break()
                add_footer(pages_count)

    # добавляем место для подписей принимающих и передающих
    document.add_paragraph("") # добавляем оступ сверху
    document.add_paragraph(_('Total count - %(co)s') % {'co': co})
    document.add_paragraph("")
    document.add_paragraph("%s        %s           ______________" % (_('Sender'), sender_full_name,))
    document.add_paragraph("")
    document.add_paragraph("%s        %s           ______________" % (_('Resipient'), recipient_full_name,))

    pages_count += 1
    add_footer(pages_count)

    document.save(file_full_name)
    
    resp_dict['error'] = '0'
    resp_dict['text']  = _('Document %(doc_number)s_%(user_id)s_0.docx generated') % { 'doc_number': m1.number, 'user_id': request.user.pk}
    #resp_dict['url'] = request.META.get('HTTP_ORIGIN') + settings.STATIC_URL + 'docx/' + docx_file_name
    resp_dict['url'] = settings.STATIC_URL + 'docx/' + docx_file_name
    
    return JsonResponse(resp_dict)


@check_ajax_auth
def generate_csv(request):
    import csv
    encoding = 'cp1251'
    def write_elem_with_group(all_items, csv_full_name):
        names_list = list()
        names_list = [str(cartridge.cart_itm_name) for cartridge in all_items]
        names_set = list(set(names_list))
        names_set.sort()
        tmp_list = []
        for name in names_set:
            tmp_list.append({name: names_list.count(name)})
        
        names_list = tmp_list
        tmp_list = names_set = None
        # формируется список из объектов-словарей, ключ - содержит имя, значение - количество повторений
        #name_list = [{'Q3312A': 3}, {'49/53A': 2}, {'505A': 5}]
        with open(csv_full_name, 'w', newline='', encoding=encoding) as csvfile:
            fieldnames = ['name', 'amount']
            writer = csv.DictWriter(csvfile, fieldnames, delimiter=';')
            writer.writerow({'name': _('Name'), 'amount': _('Amount items')})
            for name in names_list:
                key   = list(name.keys())[0]
                value = list(name.values())[0]
                writer.writerow({'name': key, 'amount': value})

    resp_dict = {}
    csv_file_name = str(int(time.time())) + '_' + str(request.user.pk) + '.csv'
    view = request.POST.get('view', '')
    gtype = request.POST.get('gtype', '')
    conf = SevercartConfigs()
    if not os.path.exists(settings.STATIC_ROOT_CSV):
        os.makedirs(settings.STATIC_ROOT_CSV)
    # Прозводим ротацию каталога csv от старых файлов
    files = filter(os.path.isfile, glob.glob(settings.STATIC_ROOT_CSV + '\*.csv'))
    files = list(files)
    files.sort(key=lambda x: os.path.getmtime(x))
    try:
        if len(files) > settings.MAX_COUNT_CSV_FILES:
            os.remove(files[0])
    except:
        pass
    csv_full_name = os.path.join(settings.STATIC_ROOT_CSV, csv_file_name)
    all_items = CartridgeItem.objects.all().order_by('pk')
    if view == 'stock':
        all_items = all_items.filter(cart_status=1).filter(departament=request.user.departament)
        if gtype == 'exp_with_group':
            write_elem_with_group(all_items, csv_full_name)
        elif gtype == 'exp_without_group':
            with open(csv_full_name, 'w', newline='', encoding=encoding) as csvfile:
                fieldnames = ['number', 'name', 'refills', 'date', 'comment']
                writer = csv.DictWriter(csvfile, fieldnames, delimiter=';')
                writer.writerow({'number': _('Number'), 
                                'name': _('Name'), 
                                'refills': _('Amount<br/>recovery'), 
                                'date': _('Date add') + ' ' + _('on stock'), 
                                'comment': _('comment')})
                for cartridje in all_items:
                    cart_date_change = cartridje.cart_date_change
                    cart_date_change = do_timezone(cart_date_change, conf.time_zone)
                    cart_date_change = cart_date_change.strftime('%d.%m.%Y %H:%M')
                    writer.writerow({'number': cartridje.cart_number, 
                                    'name': cartridje.cart_itm_name, 
                                    'refills': cartridje.cart_number_refills, 
                                    'date': cart_date_change, 
                                    'comment': cartridje.comment})
        else:
            pass
    elif view == 'use':
        try:
            root_ou   = request.user.departament
            children  = root_ou.get_family()
        except AttributeError:
            children = ''
        all_items = all_items.filter(departament__in=children).filter(cart_status=2)
        if gtype == 'exp_with_group':
            write_elem_with_group(all_items, csv_full_name)
        elif gtype == 'exp_without_group':
            with open(csv_full_name, 'w', newline='', encoding=encoding) as csvfile:
                fieldnames = ['number', 'name', 'refills', 'date', 'org', 'comment']
                writer = csv.DictWriter(csvfile, fieldnames, delimiter=';')
                writer.writerow({'number': _('Number'), 
                                'name': _('Name'), 
                                'refills': _('Amount<br/>recovery'), 
                                'date': _('Date transfe'),
                                'org': _('User'),
                                'comment': _('comment')})
                for cartridje in all_items:
                    cart_date_change = cartridje.cart_date_change
                    cart_date_change = do_timezone(cart_date_change, conf.time_zone)
                    cart_date_change = cart_date_change.strftime('%d.%m.%Y %H:%M')
                    writer.writerow({'number': cartridje.cart_number, 
                                    'name': cartridje.cart_itm_name, 
                                    'refills': cartridje.cart_number_refills, 
                                    'date': cart_date_change,
                                    'org': cartridje.departament,
                                    'comment': cartridje.comment})
        else:
            pass
    elif view == 'empty':
        all_items = all_items.filter( Q(departament=request.user.departament) & Q(cart_status=3) )
        if gtype == 'exp_with_group':
            write_elem_with_group(all_items, csv_full_name)
        elif gtype == 'exp_without_group':
            with open(csv_full_name, 'w', newline='', encoding=encoding) as csvfile:
                fieldnames = ['number', 'name', 'refills', 'date', 'comment']
                writer = csv.DictWriter(csvfile, fieldnames, delimiter=';')
                writer.writerow({'number': _('Number'), 
                                'name': _('Name'), 
                                'refills': _('Amount<br/>recovery'), 
                                'date': _('Date return'), 
                                'comment': _('comment')})
                for cartridje in all_items:
                    cart_date_change = cartridje.cart_date_change
                    cart_date_change = do_timezone(cart_date_change, conf.time_zone)
                    cart_date_change = cart_date_change.strftime('%d.%m.%Y %H:%M')
                    writer.writerow({'number': cartridje.cart_number, 
                                    'name': cartridje.cart_itm_name, 
                                    'refills': cartridje.cart_number_refills, 
                                    'date': cart_date_change,
                                    'comment': cartridje.comment})
        else:
            pass
    elif view == 'at_work':
        all_items = all_items.filter(Q(cart_status=4) & Q(departament=request.user.departament))
        if gtype == 'exp_with_group':
            write_elem_with_group(all_items, csv_full_name)
        elif gtype == 'exp_without_group':
            with open(csv_full_name, 'w', newline='', encoding=encoding) as csvfile:
                fieldnames = ['number', 'name', 'refills', 'date', 'firm', 'comment']
                writer = csv.DictWriter(csvfile, fieldnames, delimiter=';')
                writer.writerow({'number': _('Number'), 
                                'name': _('Name'), 
                                'refills': _('Amount<br/>recovery'), 
                                'date': _('Date transfer on recovery'), 
                                'firm': _('Refueller'),
                                'comment': _('comment')})
                for cartridje in all_items:
                    cart_date_change = cartridje.cart_date_change
                    cart_date_change = do_timezone(cart_date_change, conf.time_zone)
                    cart_date_change = cart_date_change.strftime('%d.%m.%Y %H:%M')
                    writer.writerow({'number': cartridje.cart_number, 
                                    'name': cartridje.cart_itm_name, 
                                    'refills': cartridje.cart_number_refills, 
                                    'date': cart_date_change,
                                    'firm': cartridje.filled_firm,
                                    'comment': cartridje.comment})
        else:
            pass
    elif view == 'basket':
        all_items = all_items.filter( (Q(cart_status=5) | Q(cart_status=6)) & Q(departament=request.user.departament) )
        if gtype == 'exp_with_group':
            write_elem_with_group(all_items, csv_full_name)
        elif gtype == 'exp_without_group':
            with open(csv_full_name, 'w', newline='', encoding=encoding) as csvfile:
                fieldnames = ['number', 'name', 'refills', 'date', 'firm', 'comment']
                writer = csv.DictWriter(csvfile, fieldnames, delimiter=';')
                writer.writerow({'number': _('Number'), 
                                'name': _('Name'), 
                                'refills': _('Amount<br/>recovery'), 
                                'date': _('Date return in basket'), 
                                'comment': _('comment')})
                for cartridje in all_items:
                    cart_date_change = cartridje.cart_date_change
                    cart_date_change = do_timezone(cart_date_change, conf.time_zone)
                    cart_date_change = cart_date_change.strftime('%d.%m.%Y %H:%M')
                    writer.writerow({'number': cartridje.cart_number, 
                                    'name': cartridje.cart_itm_name, 
                                    'refills': cartridje.cart_number_refills, 
                                    'date': cart_date_change,
                                    'comment': cartridje.comment})
        else:
            pass
    else:
        return HttpResponse(resp_dict, status_code=501)
    
    resp_dict['url'] = settings.STATIC_URL + 'csv/' + csv_file_name
    return JsonResponse(resp_dict)


@check_ajax_auth
def generate_pdf(request):
    """Генерация pdf файла с наклейками для печати.
    """
    if request.method != 'POST':
        return HttpResponse('<h1>' + _('Only use POST requests!') + '</h1>')
    
    from common.helpers import Sticker
    from service.helpers import SevercartConfigs
    resp_dict = {}
    pdf_file_name = str(int(time.time())) + '_' + str(request.user.pk) + '.pdf'
    cart_type = request.POST.get('cart_type', '')
    if not os.path.exists(settings.STATIC_ROOT_PDF):
        os.makedirs(settings.STATIC_ROOT_PDF)
    # Прозводим ротацию каталога pdf от старых файлов
    files = filter(os.path.isfile, glob.glob(settings.STATIC_ROOT_PDF + '\*.pdf'))
    files = list(files)
    files.sort(key=lambda x: os.path.getmtime(x))
    try:
        if len(files) > settings.MAX_COUNT_PDF_FILES:
            os.remove(files[0])
    except:
        pass
    pdf_full_name = os.path.join(settings.STATIC_ROOT_PDF, pdf_file_name)
    if cart_type == 'full':
        session_data = request.session.get('cumulative_list')
    elif cart_type == 'empty':
        session_data = request.session.get('empty_cart_list')
    else:
        pass

    # если сессионные данные отсутствуют, то сразу возвращаем результат
    if not session_data:
        resp_dict['url'] = ''
        return JsonResponse(resp_dict)

    session_data = json.loads(session_data)
    # формат session_data [ [name, title,  numbers=[1,2,3,4]] ... ]
    simple_cache = dict()
    list_names = CartridgeItemName.objects.all()
    for elem in list_names:
        simple_cache[elem.pk] = elem.cart_itm_name

    conf = SevercartConfigs()
    pagesize = conf.page_format
    print_bar_code = conf.print_bar_code
    pdf_doc = Sticker(file_name=pdf_full_name, pagesize=pagesize, print_bar_code=print_bar_code)
    # формируем текст для наклейки
    for elem in session_data:
        for stik in elem[2]:
            #cartridge_name = simple_cache.get(elem[0])
            cartridge_name = CartridgeItemName.objects.get(pk=elem[0])
            cartridge_name = str(cartridge_name)
            cartridge_name = cartridge_name.strip()
            cartridge_name = re.split('\s+', cartridge_name)
            # формат названия картриджа "HP RT565A"
            # если не соответсвует, то используем имя целиком.
            # если количество элементов в имени >=2 на наклейку попадёт другое имя
            if len(cartridge_name) == 1:
                cartridge_name = cartridge_name[0]
            else:
                cartridge_name = cartridge_name[-1]
            pdf_doc.add(ou_number=request.user.departament.pk, cartridge_name=cartridge_name, cartridge_number=stik)

    pdf_doc.fin()
    resp_dict['url'] = settings.STATIC_URL + 'pdf/' + pdf_file_name
    return JsonResponse(resp_dict)

@check_ajax_auth
def calculate_sum(request):
    """Подсчёт истраченных денег по заданному в списке контракту 
       на обслуживание.
    """
    ansver = dict()
    list_contracts =request.POST.getlist('service_contracts[]', [])
    try:
        list_contracts = [ int(i) for i in list_contracts ]
    except:
        list_contracts = []
    #выполняем перебор всех актов передачи для заданного
    #договра обслуживания 
    result = dict()
    for doc_id in list_contracts:
        try:
            doc = SCDoc.objects.get(pk=doc_id)
        except SCDoc.DoesNotExist:
            doc = None
        list_acts = RefillingCart.objects.filter(parent_doc=doc)
        sum = 0
        for item in list_acts:
            sum = sum + item.money

        result[doc_id] = sum 

    ansver['result'] = result
    return JsonResponse(ansver)
